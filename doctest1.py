from docx import Document
import os
import requests


class Execute:
    '''
        Execute Paragraphs KeyWords Replace
        paragraph: docx paragraph
    '''

    def __init__(self, paragraph):
        self.paragraph = paragraph


    def p_replace(self, x:int, key:str, value:str):
        '''
        paragraph replace
        The reason why you do not replace the text in a paragraph directly is that it will cause the original format to
        change. Replacing the text in runs will not cause the original format to change
        :param x:       paragraph id
        :param key:     Keywords that need to be replaced
        :param value:   The replaced keywords
        :return:
        '''
        # Gets the coordinate index values of all the characters in this paragraph [{run_index , char_index}]
        p_maps = [{"run": y, "char": z} for y, run in enumerate(self.paragraph.runs) for z, char in enumerate(list(run.text))]
        # Handle the number of times key occurs in this paragraph, and record the starting position in the list.
        # Here, while self.text.find(key) >= 0, the {"ab":"abc"} term will enter an endless loop
        # Takes a single paragraph as an independent body and gets an index list of key positions within the paragraph, or if the paragraph contains multiple keys, there are multiple index values
        k_idx = [s for s in range(len(self.paragraph.text)) if self.paragraph.text.find(key, s, len(self.paragraph.text)) == s]
        for i, start_idx in enumerate(reversed(k_idx)):       # Reverse order iteration
            end_idx = start_idx + len(key)                    # The end position of the keyword in this paragraph
            k_maps = p_maps[start_idx:end_idx]                # Map Slice List A list of dictionaries for sections that contain keywords in a paragraph
            self.r_replace(k_maps, value)
            # print(f"\t |Paragraph {x+1: >3}, object {i+1: >3} replaced successfully! | {key} ===> {value}")


    def r_replace(self, k_maps:list, value:str):
        '''
        :param k_maps: The list of indexed dictionaries containing keywords， e.g:[{"run":15, "char":3},{"run":15, "char":4},{"run":16, "char":0}]
        :param value:
        :return:
        Accept arguments, removing the characters in k_maps from back to front, leaving the first one to replace with value
        Note: Must be removed in reverse order, otherwise the list length change will cause IndedxError: string index out of range
        '''
        for i, position in enumerate(reversed(k_maps), start=1):
            y, z = position["run"], position["char"]
            run:object = self.paragraph.runs[y]         # "k_maps" may contain multiple run ids, which need to be separated
            # Pit: Instead of the replace() method, str is converted to list after a single word to prevent run.text from making an error in some cases (e.g., a single run contains a duplicate word)
            thisrun = list(run.text)
            if i < len(k_maps):
                thisrun.pop(z)          # Deleting a corresponding word
            if i == len(k_maps):        # The last iteration (first word), that is, the number of iterations is equal to the length of k_maps
                thisrun[z] = value      # Replace the word in the corresponding position with the new content
            run.text = ''.join(thisrun) # Recover



class WordReplace:
    '''
        file: Microsoft Office word file，only support .docx type file
    '''

    def __init__(self, file):
        self.docx = Document(file)

    def body_content(self, replace_dict:dict):
        print("\t☺Processing keywords in the body...")
        for key, value in replace_dict.items():
            for x, paragraph in enumerate(self.docx.paragraphs):
                Execute(paragraph).p_replace(x, key, value)
        print("\t |Body keywords in the text are replaced!")


    def body_tables(self,replace_dict:dict):
        print("\t☺Processing keywords in the body'tables...")
        for key, value in replace_dict.items():
            for table in self.docx.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for x, paragraph in enumerate(cell.paragraphs):
                            Execute(paragraph).p_replace(x, key, value)
        print("\t |Body'tables keywords in the text are replaced!")


    def header_content(self,replace_dict:dict):
        print("\t☺Processing keywords in the header'body ...")
        for key, value in replace_dict.items():
            for section in self.docx.sections:
                for x, paragraph in enumerate(section.header.paragraphs):
                    Execute(paragraph).p_replace(x, key, value)
        print("\t |Header'body keywords in the text are replaced!")


    def header_tables(self,replace_dict:dict):
        print("\t☺Processing keywords in the header'tables ...")
        for key, value in replace_dict.items():
            for section in self.docx.sections:
                for table in section.header.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for x, paragraph in enumerate(cell.paragraphs):
                                Execute(paragraph).p_replace(x, key, value)
        print("\t |Header'tables keywords in the text are replaced!")


    def footer_content(self, replace_dict:dict):
        print("\t☺Processing keywords in the footer'body ...")
        for key, value in replace_dict.items():
            for section in self.docx.sections:
                for x, paragraph in enumerate(section.footer.paragraphs):
                    Execute(paragraph).p_replace(x, key, value)
        print("\t |Footer'body keywords in the text are replaced!")


    def footer_tables(self, replace_dict:dict):
        print("\t☺Processing keywords in the footer'tables ...")
        for key, value in replace_dict.items():
            for section in self.docx.sections:
                for table in section.footer.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for x, paragraph in enumerate(cell.paragraphs):
                                Execute(paragraph).p_replace(x, key, value)
        print("\t |Footer'tables keywords in the text are replaced!")


class WordReplacer:
    def __init__(self, file):
        self.docx = Document(file)
    
    def replace_in_paragraph(self, paragraph, replace_dict):
        for idx, para in enumerate(self.docx.paragraphs):
            if para.text == paragraph:
                Execute(para).p_replace(idx, paragraph, replace_dict)
                break
            
        for table in self.docx.tables:
            for row in table.rows:
                for cell in row.cells:
                    for cell_paragraph in cell.paragraphs:
                        if cell_paragraph.text == paragraph:
                            Execute(cell_paragraph).p_replace(0, cell_paragraph.text, replace_dict)

        for section in self.docx.sections:
            for header_paragraph in section.header.paragraphs:
                if header_paragraph.text == paragraph:
                    Execute(header_paragraph).p_replace(0, header_paragraph.text, replace_dict)

            for footer_paragraph in section.footer.paragraphs:
                if footer_paragraph.text == paragraph:
                    Execute(footer_paragraph).p_replace(0, footer_paragraph.text, replace_dict)

            for header_table in section.header.tables:
                for row in header_table.rows:
                    for cell in row.cells:
                        for cell_paragraph in cell.paragraphs:
                            if cell_paragraph.text == paragraph:
                                Execute(cell_paragraph).p_replace(0, cell_paragraph.text, replace_dict)

            for footer_table in section.footer.tables:
                for row in footer_table.rows:
                    for cell in row.cells:
                        for cell_paragraph in cell.paragraphs:
                            if cell_paragraph.text == paragraph:
                                Execute(cell_paragraph).p_replace(0, cell_paragraph.text, replace_dict)

    def save(self, filepath:str):
        '''
        :param filepath: File saving path
        :return:
        '''
        print(filepath)
        self.docx.save(filepath)
        
    @staticmethod
    def docx_list(dirPath):
        '''
        :param dirPath:
        :return: List of docx files in the current directory
        '''
        fileList = []
        for roots, dirs, files in os.walk(dirPath):
            for file in files:
                if file.endswith("docx") and file[0] != "~":  # Find the docx document and exclude temporary files
                    fileRoot = os.path.join(roots, file)
                    fileList.append(fileRoot)
        print("This directory finds a total of {0} related files!".format(len(fileList)))
        return fileList
def main():
    '''
    To use: Modify the values in replace dict and filedir
    replace_dict ：key:to be replaced, value:new content
    filedir ：Directory where docx files are stored. Subdirectories are supported
    '''
    # Quan dir
    # filedir = r"C:\Users\quank\Documents\rmit\engineering science\architndesign\word_file"
    # Long dir
    # filedir = "/Users/phamlong/Desktop/RMIT/Architecture and Design/Sample Doc"
    
    # Directory where docx files are stored. Subdirectories are supported
    filedir = r"C:\Users\quank\Documents\rmit\engineering science\architndesign\word2"
    filedir2=r"C:\Users\quank\Documents\rmit\engineering science\architndesign\word2\test.docx"
    # Define the API endpoint for code generation
    api_url = "https://3c92-103-253-89-37.ngrok-free.app/generate_code?max_length=512"

    #for i, file in enumerate(WordReplacer.docx_list(filedir), start=1):
        #print(f"{i} Processing file: {file}")

        # Load the Word document
        #word_replacer = WordReplacer(filedir2)
    word_replacer = WordReplacer(filedir2)
        # Extract all paragraphs from the document
    paragraphs = [paragraph.text for paragraph in word_replacer.docx.paragraphs]
    print(paragraphs[1])
    table_texts = []
    for table in word_replacer.docx.tables:
            for row in table.rows:
                row_text = [cell.text for cell in row.cells]
                for text in row_text:
                    table_texts.append(text)
    
        # Create a list of prompts
    prompts_list = [f"Correct English grammar in the following text keep curly brackets keep it in one paragraph: {paragraph}\nHere is the corrected version: " for paragraph in paragraphs]
        # table still testing
    prompts_list_table = [f"Correct only grammar in the following text if needed do not define or add information keep it in one paragraph: {table_text}.\nHere is the corrected version: " for table_text in table_texts]
        
    all_prompts_list = prompts_list + prompts_list_table
        
        # Define API parameters
    api_params = {'prompts': all_prompts_list}
        
        # Send a GET request to the API
    response = requests.get(api_url, params=api_params)
        
        # Check the status code and response content
    if response.status_code == 200:
            corrected_paragraphs = response.json()
            
            all_text = paragraphs + table_texts

            # Replace original paragraphs with corrected paragraphs
            for i, (original, corrected) in enumerate(zip(all_text, corrected_paragraphs), start=1):
                word_replacer.replace_in_paragraph(original, corrected)
                print(f"Paragraph {i}: Replaced successfully!")
                
            # Save the document with replaced paragraphs
            output_filepath = f"document_updated.docx"
            word_replacer.save(output_filepath)
            print(f"Saved updated document to: {output_filepath}\n")
    else:
            print("Failed to retrieve corrections. Status code:", response.status_code)



if __name__ == "__main__":
    main()
    print("All complete!")